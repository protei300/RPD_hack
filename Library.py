import requests
import re
import bs4
import wget
import urllib.parse
import os
import tqdm
import secret
import asyncio, aiohttp, aiofiles
import datetime
import pandas as pd
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm
from docx.table import _Cell
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE


BASE_DIR = "D:\\RPD_TEST\\TEMP\\3"

class RPD():

    def __init__(self):
        self.login = secret.LOGIN
        self.password = secret.PASSWORD
        self.headers = {
            'Host': 'rpd.csu.ru',
            'User=Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:71.0) Gecko/20100101 Firefox/71.0',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Content-Type': 'application/x-www-form-urlencoded',
        }
        self.cookies = {
            'ASP.NET_SessionId': 'm3a1qytynxftzm2pffolfvqf',
        }

        self.get_token()

    def get_token(self):
        url = 'http://rpd.csu.ru/Auth/Login'
        params = {
            'UserName': self.login,
            'Password': self.password,

            }

        r = requests.post(url, headers=self.headers, data=params, cookies=self.cookies,
                          allow_redirects=False)
        #print (r.status_code)
        #print (r.cookies['.ASPXAUTH'])
        self.cookies['.ASPXAUTH'] = r.cookies['.ASPXAUTH']


    def make_table(self, t, df):
        # add the header rows.
        row = t.rows[0]
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        cell = row.cells[0]
        # cell.paragraphs[0].runs[0].style.font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].style.font.size = docx.shared.Pt(19)
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cell.paragraphs[0].paragraph_format.space_after = Cm(0)
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]
            t.cell(0, j).paragraphs[0].style.font.name = 'Times New Roman'
            t.cell(0, j).paragraphs[0].style.font.size = docx.shared.Pt(9)
            self.set_cell_border(
                t.cell(0, j),
                top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                end={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
            )

        # add the rest of the data frame
        for i in range(df.shape[0]):
            row = t.rows[i+1]
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            cell = row.cells[0]
            # cell.paragraphs[0].runs[0].style.font.name = 'Times New Roman'
            # cell.paragraphs[0].runs[0].style.font.size = docx.shared.Pt(19)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            cell.paragraphs[0].paragraph_format.space_after = Cm(0)
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])
                t.cell(i + 1, j).paragraphs[0].style.font.name = 'Times New Roman'
                t.cell(i + 1, j).paragraphs[0].style.font.size = docx.shared.Pt(9)
                cell = row.cells[j]
                # cell.paragraphs[0].runs[0].style.font.name = 'Times New Roman'
                # cell.paragraphs[0].runs[0].style.font.size = docx.shared.Pt(19)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Cm(0)
                self.set_cell_border(
                    t.cell(i + 1, j),
                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                )

    def to_word(self, discipline_name):



        df_main = self.df_result[self.df_result['Год'] >= 2019].\
            sort_values('Год', ascending=False)
        del df_main['Год']

        df_support = self.df_result[(self.df_result['Год'] >= 2016) & (self.df_result['Год']<2019)].\
            sort_values('Год', ascending=False)
        del df_support['Год']

        doc = docx.Document()
        # add a table to the end and create a reference variable
        # extra row is so we can add the header row


        run = doc.add_paragraph().add_run('Основная литература')

        run.font.name = 'Times New Roman'
        run.font.size = docx.shared.Pt(14)
        run.font.bold = True

        t = doc.add_table(df_main.shape[0] + 1, df_main.shape[1])
        self.make_table(t, df_main)

        paragraph = doc.add_paragraph(' ')
        run = doc.add_paragraph().add_run('Дополнительная литература')
        run.font.name = 'Times New Roman'
        run.font.size = docx.shared.Pt(14)
        run.font.bold = True

        t = doc.add_table(df_support.shape[0] + 1, df_support.shape[1])
        self.make_table(t, df_support)

        # add the header rows.


        # save the doc
        doc.save(os.path.join(BASE_DIR,f"{discipline_name}_{self.request}.docx"))


    def get_books(self, request):
        self.cookies['authors'] = 'true'
        self.cookies['litName'] = 'true'
        self.cookies['imprintDate'] = 'true'
        self.cookies['publishing'] = 'true'
        #self.cookies['edition'] = 'true'
        self.cookies['adress'] = 'true'

        self.request = request

        data = {
            'LitName': request,
            'ImprintDate': '',
            'IsPeriodic': 'false',
            'command': 'search',
            'RupRowId': 1665012,
            'RpdId': 31315,
            'WorkId': 924758,
            'LitType': 1,
        }

        url = "http://rpd.csu.ru/LitManager/SearchLit"

        res = requests.post(url=url, data=data, headers=self.headers, cookies=self.cookies)

        soup = bs4.BeautifulSoup(res.text, features='lxml', parser= 'html.parser')
        table = soup.find_all('table', id='litGrid_DXMainTable')[0]
        #print(type(table))
        df = pd.read_html(str(table))[0]
        df.dropna(inplace=True, how='all')
        df.dropna(inplace=True, how='all', axis=1)
        df.columns = ['Название', 'Автор', 'Издательство', 'Год', 'Ссылка']
        df['Год'] = df['Год'].astype('str')
        df['Год'] = df['Год'].str.replace(']','')
        df = df[~df['Год'].isin(['Б. г.', 'б.г.', '[б. г.]', '[б. г.'])]

        df = df[~df['Год'].str.contains('-')]
        df.dropna(inplace=True)
        df['Год'] = df['Год'].astype('float')
        df['Год'] = df['Год'].astype('int16')
        df['Год'] = df['Год'].astype('str')
        df_result = pd.DataFrame(columns = ['Столбец1', 'Столбец2', 'Столбец3'],
                                 index = df.index)

        df_result['Столбец1'] = df['Автор']
        df_result['Столбец2'] = df['Название'] + '\n' + '(' + df['Ссылка'] + ')' + '\n' + f"Дата обращения: {datetime.datetime.now().strftime('%Y-%m-%d')}"
        df_result['Столбец3'] = df['Издательство'] + ', ' + df['Год']
        df['Год'] = df['Год'].astype('int16')
        df_result['Год'] = df['Год']
        self.df_result = df_result



    def set_cell_border(self, cell: _Cell, **kwargs):
        """
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


rpd = RPD()
for word in ['Компьютерные игры']:
    rpd.get_books(word)
    rpd.to_word('разное')


