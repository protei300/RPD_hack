'''
Модуль по удалению и преобразованию колонтитулов

'''


import re
import requests
import urllib.parse
import os
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Mm
import copy
from tqdm import tqdm

BASE_DIR = "D:\\RPD_TEST\\TEMP"
RPD_DIR = "D:\\RPD_TEST\\RPD_LOADED"
RPD_MODIFIED_DIR = "D:\\RPD_TEST\\RPD8"
import secret

PARAGRAFS_SER = [
    'Титульный лист',
    'Содержание',
    '1. Цели освоения дисциплины',
    '2. Место дисциплины в структуре ОПОП',
    #'2.1 Требования к предварительной подготовке обучающегося:',
    #'2.2 Дисциплины и практики, для которых освоение данной дисциплины (модуля) необходимо как предшествующее:',
    '3. Компетенции обучающегося, формируемые в результате освоения дисциплины (модуля)',
    '4. Объем дисциплины (модуля)',
    '5. Структура и содержание дисциплины (модуля)',
    '6. Фонд оценочных средств',
    #'6.1. Перечень видов оценочных средств',
    #'6.2. Типовые контрольные задания и иные материалы для текущей аттестации',
    #'6.3. Типовые контрольные вопросы и задания для промежуточной аттестации',
    #'6.4. Критерии оценивания',
    '7. Учебно-методическое и информационное обеспечение дисциплины (модуля)',
    #'7.1. Рекомендуемая литература',
    #'7.1.1. Основная литература',
    #'7.1.2. Дополнительная литература',
    #'7.2. Перечень ресурсов информационно-телекоммуникационной сети "Интернет"',
    #'7.3 Перечень информационных технологий',
    #'7.3.1 Программное обеспечение',
    #'7.3.2 Информационно-справочные системы',
    '8. Материально-техническое обеспечение дисциплины (модуля)',
    '9. Методические указания для обучающихся по освоению дисциплины (модуля)',
    '10. Специальные условия освоения дисциплины обучающимися с инвалидностью и ограниченными возможностями здоровья'
]

FOOTER = '© ФГБОУ ВО «ЧелГУ»'

class Docx():
    def __init__(self, filepath, filename, path_to_save, filename_to_save):
        #self.filepath = filepath
        self.filename = filename
        self.path_to_save = path_to_save
        self.filename_to_save = filename_to_save
        self.path_to_files = filepath


    def create_footer(self, footer):
        footer.paragraphs[0].text = FOOTER
        footer.paragraphs[0].style.font.name = 'Times New Roman'
        footer.paragraphs[0].style.font.size = docx.shared.Pt(9.5)

    def create_header(self, document):
        section = document.sections[0]
        header = section.header

        first_header = docx.Document(os.path.join(BASE_DIR, 'Insertions', 'first_header.docx'))
        cell = first_header.tables[1].rows[0].cells[0]
        cell.text = self.ESCAPE_SENTENCE[1].replace('\n', '')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].style.font.name = 'Times New Roman'
        cell.paragraphs[0].style.font.size = docx.shared.Pt(9)

        header_table = copy.deepcopy(first_header.tables[1]._tbl)
        header.paragraphs[0]._p.addnext(header_table)
        self.delete_paragraph(header.paragraphs[0])
        #document.save(os.path.join(RPD_MODIFIED_DIR, self.filepath, self.filename_to_save))

        #document = docx.Document(os.path.join(RPD_MODIFIED_DIR, self.filepath, self.filename_to_save))
        return document
        #section = document.sections[0]


    def create_first_page(self, document):

        first_header = docx.Document(os.path.join(BASE_DIR,'Insertions', 'first_header.docx'))

        cell = first_header.tables[0].rows[1].cells[0]
        cell.text = self.ESCAPE_SENTENCE[1].replace('\n','')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].style.font.name = 'Times New Roman'
        cell.paragraphs[0].style.font.size = docx.shared.Pt(9)

        first_header_table = copy.deepcopy(first_header.tables[0]._tbl)

        section = document.sections[0]
        header = section.first_page_header
        header.paragraphs[0]._p.addnext(first_header_table)
        self.delete_paragraph(header.paragraphs[0])
        #document.save(os.path.join(RPD_MODIFIED_DIR, self.filepath, self.filename_to_save))

        #document = docx.Document(os.path.join(RPD_MODIFIED_DIR, self.filepath, self.filename_to_save))
        #section = document.sections[0]

        #section.first_page_header.paragraphs[0].text = ''
        cell = section.first_page_header.tables[0].cell(0,1)
        cell._element.clear_content()
        cell.add_paragraph().add_run().add_picture(os.path.join(BASE_DIR,'Insertions', 'csu.png'),
                                               width=Cm(1.75))

        return document




    def generate_docx(self):
        document = docx.Document(os.path.join(self.path_to_files, self.filename))
        self.ESCAPE_SENTENCE = (FOOTER, document.tables[0].rows[1].cells[0].text)
        section = document.sections[0]

        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(20)
        section.right_margin = Mm(10)
        section.top_margin = Mm(1)
        section.bottom_margin = Mm(9.5)
        section.header_distance = Mm(8)
        #section.footer_distance = Mm(12.7)
        section.different_first_page_header_footer = True

        latent_styles = document.styles.latent_styles
        latent_styles.add_latent_style('Table Grid')

        self.create_footer(section.footer)
        self.create_footer(section.first_page_footer)

        document = self.create_first_page(document)
        document = self.create_header(document)

        ##### удаляем колонтитул на 1 странице ###############
        for row in document.tables[0].rows[:2]:
            self.remove_row(document.tables[0], row)

        ##### Удаляем остальные колонтитулы и футеры #########
        for table in document.tables:
            self.clear_headers(table)

        #print(self.paragrafs['Титульный лист'])
        #print(len(document.paragraphs))
        for para in document.paragraphs[2:]:
            self.delete_paragraph(para)
        document.save(os.path.join(self.path_to_save, self.filename_to_save))

    def print_paragrafs(self):
        for key,value in self.paragrafs.items():
            print(f"{key}: {value}")

    @staticmethod
    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None



    def clear_headers(self, table):
        for row in table.rows:
            #print(row.cells[0].text)
            if row.cells[0].text in self.ESCAPE_SENTENCE or self.ESCAPE_SENTENCE[0] in row.cells[0].text:
                #print(f"DELETING {row.cells[0].text}")
                self.remove_row(table, row)


def clean_documents(rpd_dir, rpd_modified_dir):

    for dir_l_1 in os.listdir(rpd_dir):
        print(f"Proceeding to {dir_l_1} subdir")
        for dir_l_2 in os.listdir(os.path.join(rpd_dir,dir_l_1)):
            print(f"Proceeding to {dir_l_2} subdir")
            middle_path = os.path.join(dir_l_1, dir_l_2)
            path = os.path.join(rpd_dir, middle_path)
            path_to_save = os.path.join(rpd_modified_dir, middle_path)
            if not os.path.isdir(path_to_save):
                os.makedirs(path_to_save)

            for file in tqdm(os.listdir(path)):
                filename_to_save = file.split('_')[-1]

                try:
                    docx_doc = Docx(
                                    path,
                                    file,
                                    path_to_save,
                                    filename_to_save,
                    )
                    docx_doc.generate_docx()
                except Exception as e:
                    print(f"Error reading - {file}. Error - {e}")
                    with open(os.path.join (rpd_modified_dir, middle_path, filename_to_save), 'w') as f:
                        pass


def clean_practics(path):
    for file in tqdm(os.listdir(path)):
        filename_to_save = f'!{file}'
        try:
            docx_doc = Docx(
                path,
                file,
                filename_to_save,
            )
            docx_doc.generate_docx()
        except Exception as e:
            print(f"Error reading - {file}. Error - {e}")

def main():
    #clean_practics(r"D:\RPD_TEST\Практики")
    clean_documents(rpd_dir=r'D:\RPD_TEST\RPD_TO_MOD',
                    rpd_modified_dir=RPD_MODIFIED_DIR)


if __name__ == '__main__':
    main()


