import os
import re
import bs4
import docx
import tempfile
import base64
from docx.shared import Inches

TEMPDIR = tempfile.mkdtemp()
MAX_INCHES_WIDTH = 3  #Максимальная ширина картинки
WORKING_DIR = "D:\\RPD_TEST"  #Рабочая директория
WORKING_DIR_XML = os.path.join(WORKING_DIR,"XML")   # Место хранения XML файлов
WORKING_DIR_DOCX = os.path.join(WORKING_DIR, "DOCX") # Место хранения Docx Файлов


class Question:
    def __init__(self, **kwargs):

        self.name = kwargs['name']
        self.fullquestion = kwargs['full_question']
        self.variants = kwargs['variants']
        for i, variant in enumerate(self.variants):
            self.variants[i]['text']=variant['text'].replace('\n', '')

    def get_variants(self):
        return self.variants

    def __str__(self):
        #print (self.fullquestion)
        #print ([variant['text'] for variant in self.variants])
        variant_string = f'{self.fullquestion}\n'
        for variant in self.variants:
            variant_string += variant['text'] + '\n'

        return str(variant_string)


def make_docx(questions, docx_name):
    '''
    Функция создает docx файл в виде таблицы, заполняя его данными из словаря вопросов
    :param questions:  словарь вопросов (поддерживает 3 типа вопросов)
    :param docx_name: имя файла docx для сохранения
    :return:
    '''

    document = docx.Document()

    table = document.add_table(rows=1, cols=3)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    table.autofit = True
    heading_cells = table.rows[0].cells
    table.style = 'Table Grid'

    #Заголовок таблицы
    heading_cells[0].text = '№ п/п'
    heading_cells[1].text = 'Формулировка вопроса'
    heading_cells[2].text = 'Варианты ответов'

    quant = 1
    for key in questions:
        for ques in questions[key]:
            cells = table.add_row().cells

            #Добавляем индекс вопроса
            cells[0].text = f"{quant}."

            #Обрабатываем текст вопроса
            fullquestion = ques.fullquestion
            if re.findall(r'(\d+.png)', fullquestion) != []:
                filename = re.findall(r'(\d+.png)', fullquestion)[0]
                run = cells[1].paragraphs[0].add_run()
                splitted_full_question = fullquestion.split(filename)[0]
                run.add_text(splitted_full_question)
                pic_run = run.add_picture(os.path.join(TEMPDIR, filename))
                if pic_run.width.inches > MAX_INCHES_WIDTH:
                    pic_run.height = Inches(pic_run.height.inches*MAX_INCHES_WIDTH/pic_run.width.inches)
                    pic_run.width = Inches(MAX_INCHES_WIDTH)

                #print(pic_run.width.inches, pic_run.height.inches)
                run.add_break(break_type=6)
            else:
                cells[1].text = fullquestion

            #Обрабатываем вопросы
            variants = ques.get_variants()
            if key != 'matching':
                for variant in variants:
                    splitted_text = variant['text'].split(" - ")[1]
                    run = cells[2].paragraphs[0].add_run(f"{splitted_text}\n")
                    if float(variant['text'].split(" - ")[0])> 0:
                        run.bold=True
            else:
                for variant in variants:
                    #print (variant['text'])
                    if re.findall(r'^(\d+.png)',variant['text']) != []:
                        filename = re.findall(r'^(\d+.png)',variant['text'])[0]
                        run = cells[2].paragraphs[0].add_run()
                        run.add_picture(os.path.join(TEMPDIR, filename))
                        cutted_variant = variant['text'].split(' - ')[1]
                        run.add_text(f" - {cutted_variant}")
                        run.add_break(break_type=6)

                    else:
                        cells[2].paragraphs[0].add_run(f"{variant['text']}\n\n")
            quant += 1


    document.save(os.path.join(WORKING_DIR_DOCX, docx_name))



def create_tests(xml_name):
    '''
    Функция читает данные из xml файла, получает словарь questions
    :param xml_name: имя файла xml для чтения
    :return: словарь вопросов
    '''
    if  not os.path.exists(WORKING_DIR_DOCX):
            os.mkdir(WORKING_DIR_DOCX)

    with open(os.path.join(WORKING_DIR_XML, xml_name) , 'r', encoding="utf-8") as f:
        file = f.read()

    questions = {'multichoice': [],
                 'matching':[],
                 'truefalse':[]
                 }

    filename = 1
    soup = bs4.BeautifulSoup(file, features='html.parser')
    for ques in soup.find_all('question'):
        #print ('\n')

        if ques['type'] == 'matching':

            answers = []
            for subques in ques.find_all('subquestion'):
                if subques.file == None:
                    answers.append ( {
                        'text': f"{subques.find('text').get_text().replace('<p>','').replace('</p>','')} - {subques.answer.get_text()}"})
                else:
                    answers.append({
                        'text': f"{filename}.png - {subques.answer.get_text()}",
                    })

                    with open(os.path.join(TEMPDIR, f"{filename}.png"), "wb") as f:
                        f.write(base64.b64decode((subques.file.get_text())))
                    #print (f"{filename}.png")
                    filename += 1

            q = Question(name =  ques.questiontext.get_text().replace('<p>','').replace('</p>',''),
                       full_question = ques.questiontext.get_text().replace('<p>','').replace('</p>',''),
                       variants =  answers)
            questions['matching'].append(q)


        elif ques['type'] == 'multichoice':
            answers = []
            #print (ques.questiontext.find(text=lambda tag: isinstance(tag, bs4.CData)))
            if ques.questiontext.find(text=lambda tag: isinstance(tag, bs4.CData)) != None:
                #print (ques.questiontext.find(text=lambda tag: isinstance(tag, bs4.CData)))
                temp_soup = bs4.BeautifulSoup(ques.questiontext.find(text=lambda tag: isinstance(tag, bs4.CData)), features='html.parser')
                full_question = temp_soup.text
                if ques.questiontext.file != None:
                    with open(os.path.join(TEMPDIR, f"{filename}.png"), "wb") as f:
                        f.write(base64.b64decode((ques.questiontext.file.get_text())))

                    full_question+= f' {filename}.png'
                    filename += 1
                #print (full_question)

            else:
                full_question = ques.questiontext.get_text()

            for subques in ques.find_all('answer'):

                answers.append({
                    'text': f"{subques['fraction']} - {subques.find('text').string.replace('<p>','').replace('</p>','')}",
                    'file': None
                })
            q = Question(name  = full_question, full_question = full_question, variants=answers)
            questions['multichoice'].append(q)



        elif ques['type'] == 'truefalse':
            answers = []
            if ques.questiontext.find(text=lambda tag: isinstance(tag, bs4.CData)) != None:
                # print (ques.questiontext.find(text=lambda tag: isinstance(tag, bs4.CData)))
                temp_soup = bs4.BeautifulSoup(ques.questiontext.find(text=lambda tag: isinstance(tag, bs4.CData)),
                                              features='html.parser')
                full_question = temp_soup.text
                #print(full_question)

            else:
                full_question = ques.questiontext.get_text()
            for subques in ques.find_all('answer'):
                answers.append({
                    'text': f"{subques['fraction']} - {'Верно' if subques.find('text').get_text() == 'true' else 'Неверно'}",
                    'file': None
                })
            q = Question(name=full_question, full_question=full_question, variants=answers)
            questions['truefalse'].append(q)
    return questions




def main():
    for xml_file in os.listdir(WORKING_DIR_XML):
        questions = create_tests(xml_file)
        docx_name = re.findall(r'^вопросы-([\w\s.-]+) \(', xml_file)[0]
        make_docx(questions, f"{docx_name}.docx")








if __name__ == '__main__':


    main()




'''with open("D:\\RPD_TEST\\foo.png","wb") as f:
    f.write(base64.b64decode('/9j/4AAQSkZJRgABAQEAYABgAAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAAiABQDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD9/KoeFP8AkVtN/wCvWL/0AV85/wDBTHxx/Z3wp0n4f2E9yuo/Eq8NjchZ33R6VCBJfN16Ohjtvrdj0rnf+CY/7K/hnwhDrPxK0rw5o/h9PEUP9l6ONNs0svtFhHJue5bywu4TTL8hJZTFBC6481hXxj4vU+Jf9XMNRc3CmqlSd7KF3aMWrO7lutVoez/ZHLlv9o1J2vLljG2srbv0R9e0UiKEUAZwBjk5NFfZnjHwT/wUf+GvxG8PfEfV/iJfQ+Ctb8MS2Nh4V0S0XX7zTNVzPL+9SKJLWRRLJLJuLidcR2yMSuw47X9gH9mjW9Mm0rxrb+KfG1p4Us7H7L4f0fUvE2rahZX0DRhFuDbT3RiWBUAEChOR+8HBjr1T47fsw3P7Rvx28Han4mlMvgDwXBLeR6EjKTrGpy7ot9zk7fIjgyoQE+Z9olV8ICsnrfh55bLQLGFrWfMVvGh+6vRQOhII/EA1+c5dwDhaPEmI4hTnFz+zzz5ZSt8co83Lp8MFayV3u1b6HEZ9Vnl1PL2k0uvKrpdk7X13b67GjEHWJRIys4A3FV2gnuQMnH5miiJzJGCUZCf4WxkflkUV+jHzw6iiigAooooA/9k='))'''

