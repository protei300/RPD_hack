import docx
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas
import re
import tqdm


PARAGRAFS_SER = [
    '1. Цели освоения дисциплины',
    '2. Место дисциплины в структуре ОПОП',
    #'2.1 Требования к предварительной подготовке обучающегося:',
    #'2.2 Дисциплины и практики, для которых освоение данной дисциплины (модуля) необходимо как предшествующее:',
    '3. Компетенции обучающегося, формируемые в результате освоения дисциплины (модуля)',
    '4. Объем дисциплины (модуля)',
    '5. Структура и содержание дисциплины (модуля)',
    '6. Фонд оценочных средств',
    #'6.1. Перечень видов оценочных средств',
    #'6.2. Типовые контрольные задания и иные материалы для текущей аттестации',
    #'6.3. Типовые контрольные вопросы и задания для промежуточной аттестации',
    #'6.4. Критерии оценивания',
    '7. Учебно-методическое и информационное обеспечение дисциплины (модуля)',
    #'7.1. Рекомендуемая литература',
    #'7.1.1. Основная литература',
    #'7.1.2. Дополнительная литература',
    #'7.2. Перечень ресурсов информационно-телекоммуникационной сети "Интернет"',
    #'7.3 Перечень информационных технологий',
    #'7.3.1 Программное обеспечение',
    #'7.3.2 Информационно-справочные системы',
    '8. Материально-техническое обеспечение дисциплины (модуля)',
    '9. Методические указания для обучающихся по освоению дисциплины (модуля)',
    '10. Специальные условия освоения дисциплины обучающимися с инвалидностью и ограниченными возможностями здоровья'
]



PARAGRAFS = [

    '8. Материально-техническое обеспечение дисциплины (модуля)',
    '9. Методические указания для обучающихся по освоению дисциплины (модуля)',
    '10. Специальные условия освоения дисциплины обучающимися с инвалидностью и ограниченными возможностями здоровья'
]

learn_form = re.compile(r'^\d{4}_\((очная|очно-заочная|заочная)\)_')

DATA_DIR = "F:\\RPD_TEST"
WORKING_DIR = "F:\\RPD_TEST\\SOURCE"
SAVE_DIR = "F:\\RPD_TEST\\RESULT"

def serialyze(paragrafs_content, keys, table_content):
    current_para_key = None
    next_para_key = 0
    for i,content in enumerate(table_content):
        if keys[next_para_key].lower() in content.lower():
            current_para_key = next_para_key
            if next_para_key + 1 < len(keys):
                next_para_key += 1
        elif current_para_key != None and keys[next_para_key].lower() not in content.lower():
            if content != '' and 'ФГБОУ ВО «ЧелГУ»'.lower() not in content.lower():
                paragrafs_content[keys[current_para_key]].append(content)


def main():

    df = pandas.read_excel(os.path.join(DATA_DIR,'Данные.xlsx'),index_col='Параграф')
    text_to_add_o = df['Очка'].to_dict()
    text_to_add_o_k = df['Очка_Курс'].to_dict()
    text_to_add_z = df['Заочка'].to_dict()
    text_to_add_z_k = df['Заочка_Курс'].to_dict()



    files = os.listdir(WORKING_DIR)
    for file in tqdm.tqdm(files):

        filename = os.path.splitext(file)[0]
        form = learn_form.findall(filename)[0]
        print (filename)
        content = docx.Document(os.path.join(WORKING_DIR,file))

        # Создаем список разделов в документе
        paragrafs_content = dict.fromkeys(PARAGRAFS_SER, 0)
        for key in paragrafs_content.keys():
            paragrafs_content[key] = []

        table_content = []
        # Делаем словарь уровня 0



        i = 0
        modification = False
        serialyzed = False
        cursov = False
        for row in content.tables[-1].rows:

            if len(row.cells)>0:

                if not modification:
                    for cell in row.cells[:4]:
                        if PARAGRAFS[i].lower() in cell.text.lower():
                            print (PARAGRAFS[i])
                            modification = True
                            if not serialyzed:
                                serialyze(paragrafs_content, PARAGRAFS_SER, table_content)
                                serialyzed = True
                                for paragraf_content in paragrafs_content[PARAGRAFS_SER[3]]:
                                    if 'курсовые работы' in paragraf_content:
                                        cursov = True
                                        break
                            break

                else:
                    if form.lower() == 'очная' and cursov == False:
                        row.cells[0].text = text_to_add_o[list(text_to_add_o.keys())[i]]
                    elif form.lower() == 'очная' and cursov == True:
                        row.cells[0].text = text_to_add_o_k[list(text_to_add_o_k.keys())[i]]
                    else:
                        row.cells[0].text = text_to_add_z[list(text_to_add_z.keys())[i]]
                    para = row.cells[0].paragraphs[0]
                    para.style.font.name = 'Times New Roman'
                    para.style.font.size = docx.shared.Pt(9.5)
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    modification = False
                    if i<len(PARAGRAFS)-1:
                        i+=1
                for cell in row.cells:
                    if cell.text!='':
                        table_content.append(cell.text)

        #serialyze(paragrafs_content, PARAGRAFS_SER, table_content)
        #print(paragrafs_content[PARAGRAFS_SER[3]])
        #cursov = False
        '''for paragraf_content in paragrafs_content[PARAGRAFS_SER[3]]:
            if 'курсовые работы' in paragraf_content:
                cursov = True
                break'''

        print (cursov)


        #filename = os.path.splitext(files[0])[0]
        content.save(os.path.join(SAVE_DIR,file))





if __name__ == '__main__':

    main()
