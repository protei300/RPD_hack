import re
import requests
import urllib.parse
import os
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm
from docx.table import _Cell
import copy
from tqdm import tqdm
import pandas as pd



PARAGRAFS_SER = [
    'Титульный лист',
    'Содержание',
    '1. Цели освоения дисциплины',
    '2. Место дисциплины в структуре ОПОП',
    #'2.1 Требования к предварительной подготовке обучающегося:',
    #'2.2 Дисциплины и практики, для которых освоение данной дисциплины (модуля) необходимо как предшествующее:',
    '3. Компетенции обучающегося, формируемые в результате освоения дисциплины (модуля)',
    '4. Объем дисциплины (модуля)',
    '5. Структура и содержание дисциплины (модуля)',
    '6. Фонд оценочных средств',
    #'6.1. Перечень видов оценочных средств',
    #'6.2. Типовые контрольные задания и иные материалы для текущей аттестации',
    #'6.3. Типовые контрольные вопросы и задания для промежуточной аттестации',
    #'6.4. Критерии оценивания',
    '7. Учебно-методическое и информационное обеспечение дисциплины (модуля)',
    #'7.1. Рекомендуемая литература',
    #'7.1.1. Основная литература',
    #'7.1.2. Дополнительная литература',
    #'7.2. Перечень ресурсов информационно-телекоммуникационной сети "Интернет"',
    #'7.3 Перечень информационных технологий',
    #'7.3.1 Программное обеспечение',
    #'7.3.2 Информационно-справочные системы',
    '8. Материально-техническое обеспечение дисциплины (модуля)',
    '9. Методические указания для обучающихся по освоению дисциплины (модуля)',
    '10. Специальные условия освоения дисциплины обучающимися с инвалидностью и ограниченными возможностями здоровья'
]




SECTIONS = {
    "1": '1. Цели освоения дисциплины'.upper(),
    "2": '2. Место дисциплины в структуре ОПОП'.upper(),
    "3": '3. Компетенции обучающегося, формируемые в результате освоения дисциплины'.upper(),
    "4": '4. Объем дисциплины'.upper(),
    "7.3.2": r'(7.3.2 Информационно-справочные системы|7.3.2 Профессиональные базы данных и информационно-справочные системы)',
    "8": '8. Материально-техническое обеспечение дисциплины'.upper(),
    "9": '9. Методические указания для обучающихся по освоению дисциплины'.upper(),
    "10": '10. Специальные условия освоения дисциплины обучающимися с'.upper(),

}

COMPETENCE_REGEX = re.compile(r'(ОПК-\d{1,2}|УК-\d{1,2}|ПК-\d{1,2}):')
LEARNING_CODE_REGEX = re.compile(r'(\d{2}\.\d{2}\.\d{2})')
LEARNING_YEAR_REGEX = re.compile(r'Челябинск (\d{4}) г.')
LEARNING_FORM_REGEX = re.compile(r'(очная|заочная|очно-заочная)')

BASE_DIR = "D:\\RPD_TEST\\"

class RPD_CHAPTERS:

    def __init__(self):
        self.learning_plans = pd.read_excel(os.path.join(BASE_DIR,'data','компетенции.xlsx'),
                                   sheet_name=None)

        for code, plan in self.learning_plans.items():
            plan['Компетенция'] = plan['Индикаторы'].str.extract(r'(ПК-\d{1,2}|ОПК-\d{1,2}|УК-\d{1,2})')

        self.chapters = pd.read_excel(os.path.join(BASE_DIR,'data','разделы.xlsx'), sheet_name=None)


    def get_learning_plan(self, plan_code):
        if plan_code in self.learning_plans.keys():
            return 0, self.learning_plans[plan_code]
        else:
            return 1, None


    def get_chapter(self, chapter_code):
        return self.chapters[chapter_code]


class Sections_changer():

    rpd_chapters = RPD_CHAPTERS()

    def __init__(self, path_to_read, path_to_save, filename):
        self.filename = filename
        self.path_to_save =  path_to_save
        self.document = docx.Document(os.path.join(path_to_read, self.filename))
        self.docx_map = dict.fromkeys(SECTIONS.keys())
        self.competence = {}
        self.learning_plan_info = dict.fromkeys(['learning_code', 'learning_form', 'year', 'kurs'])
        self.map_document()


    def map_document(self):
        '''
        Функция разметки docx файла.
        Формирует словарь с указанием номера таблицы и номера строки, где начинается соответствующий раздел из SECTIONS
        Так же формирует список компетенций, который есть в файле
        :return:
        '''
        key_number = 0
        keys = list(SECTIONS.keys())

        # Ищем код направления
        for row in self.document.tables[0].rows:
            learning_code = LEARNING_CODE_REGEX.findall(row.cells[0].text)
            learning_year = LEARNING_YEAR_REGEX.findall(row.cells[0].text)
            learning_form = LEARNING_FORM_REGEX.findall(row.cells[0].text)
            if len(learning_code) != 0:
                self.learning_plan_info['learning_code'] = learning_code[0]
            elif len(learning_year) != 0:
                self.learning_plan_info['year'] = int(learning_year[0])
            elif len(learning_form) != 0:
                if learning_form[0] == 'очная':
                    self.learning_plan_info['learning_form'] = True
                else:
                    self.learning_plan_info['learning_form'] = False


        # Ищем номера таблиц и строк где находятся значимые разделы
        for table_num, table in enumerate(self.document.tables[2:]):
            table_num += 2 #делаем поправку на то, что 2 таблицы первые нам не нужны
            for row_num, row in enumerate(table.rows):
                if len(row.cells) == 0:
                    continue
                text_to_search = ''
                for cell in row.cells[:3]:
                    text_to_search += cell.text


                if re.search(SECTIONS[keys[key_number]], text_to_search) and text_to_search != '':
                    self.docx_map[keys[key_number]] = (table_num, row_num)
                    if key_number < len(keys)-1:
                        key_number += 1
                    else:
                        break
                    #print(f'Searching key {SECTIONS[keys[key_number]]}')
                else:
                    compet = self.find_competence(row.cells[0].text)
                    if compet:
                        self.competence[compet] = ''
                #print(' ')
                #print(text_to_search)

        code, indicators = self.rpd_chapters.get_learning_plan(self.learning_plan_info['learning_code'])
        if code == 0:
            indicators = indicators[indicators['Компетенция'].isin(list(self.competence.keys()))]

            for i,row in indicators.iterrows():
                self.competence[row['Компетенция']] += f"{row['Индикаторы']}\n"

        try:
            table_num, row_num = self.docx_map['4']

            for cell in self.document.tables[table_num].rows[row_num+2].cells:
                if len(re.findall('курсовые работы', cell.text)) != 0:
                    self.learning_plan_info['kurs'] = True
        except:
            pass

        #print(self.docx_map)



        #print("Разметка документа завершена")
        #print("Найдены следующие компетенции")
        #print(list(self.competence.keys()))
        #print("Так же найдены такие свойства")
        #print(self.learning_plan_info.values())

    def find_competence(self, text):
        '''
        Функция поиска номера компенетции в строке
        :param text:
        :return:
        '''
        res = COMPETENCE_REGEX.findall(text)
        if len(res) != 0:
            return res[0]
        else:
            return False

    def stop_words(self,text):
        '''
        Проверка текста на ненужные предложения, которые должны быть исключены
        :param text:
        :return:
        '''
        STOP_WORDS = (
            '© ФГБОУ ВО «ЧелГУ»',
            'Рабочая программа дисциплины',
            '3.',
            'В результате освоения дисциплины',

                                  )

        for stop_word in STOP_WORDS:
            if stop_word in text:
                return True
        if text == '':
            return True
        return False


    def read_three(self):
        '''
        Ищет описание компетенций в 3 разделе
        :return:
        '''
        table_from, row_from = self.docx_map['3']
        competence_text = ''
        if self.docx_map['4'] is None:
            table_till = len(self.document.tables)
            row_till = len(self.document.tables[-1].rows)
        else:
            table_till, row_till = self.docx_map['4']
            table_till += 1

        if table_from == table_till - 1:
            for row_num in range(row_from+1, row_till):
                print(self.document.tables[table_from].rows[row_num].cells[0].text)
                if self.document.tables[table_from].rows[row_num].cells[0].text == '':
                    competence_text += self.document.tables[table_from].rows[row_num].cells[1].text + '\n'
                else:
                    competence_text += self.document.tables[table_from].rows[row_num].cells[0].text + '\n'
        else:
            for table_num in range(table_from, table_till):
                print(table_num, len(self.document.tables[table_num].rows))
                if table_num == table_from:
                    for row_num in range(row_from+1, len(self.document.tables[table_num].rows)):
                        if len(self.document.tables[table_num].rows[row_num].cells) == 0: continue
                        if not self.stop_words(self.document.tables[table_num].rows[row_num].cells[0].text):
                            competence_text += self.document.tables[table_num].rows[row_num].cells[0].text + '\n'
                elif table_from < table_num < table_till:
                    for row_num in range(len(self.document.tables[table_num].rows)):
                        if len(self.document.tables[table_num].rows[row_num].cells) == 0: continue
                        if not self.stop_words(self.document.tables[table_num].rows[row_num].cells[0].text):
                            competence_text += self.document.tables[table_num].rows[row_num].cells[0].text + '\n'
                else:
                    for row_num in range(row_till):
                        if len(self.document.tables[table_num].rows[row_num].cells) == 0: continue
                        if not self.stop_words(self.document.tables[table_num].rows[row_num].cells[0].text):
                            competence_text += self.document.tables[table_num].rows[row_num].cells[0].text + '\n'
        print(competence_text)
        competence_text = competence_text.split('\n')[:-1]

        competence_descriptor, competence_keys = dict.fromkeys(self.competence.keys()), list(self.competence.keys())
        competence_number = 0

        for line_number, line in enumerate(competence_text):
            if competence_keys[competence_number] in line:
                competence_descriptor[competence_keys[competence_number]] = line_number
                if competence_number == len(competence_keys)-1: continue
                competence_number += 1

        print(competence_descriptor)










    def change_one(self):
        '''
        Добавляем индикаторы в 1 раздел, если план 3++
        :return:
        '''
        if self.learning_plan_info['year'] >= 2019:
            table_num, row_num = self.docx_map['2']
            self.add_row(self.document.tables[table_num].rows[row_num-2],
                         self.document.tables[table_num].rows[row_num-2])

            cell = self.document.tables[table_num].rows[row_num-1].cells[0]

            self.document.tables[table_num].rows[row_num - 1].height_rule = WD_ROW_HEIGHT_RULE.AUTO
            cell.paragraphs[0].style.font.name = 'Times New Roman'
            cell.paragraphs[0].style.font.size = docx.shared.Pt(9)
            text = 'Результаты обучения по дисциплине направлены на достижение индикаторов:\n'
            keys = sorted(self.competence.keys())
            for key in keys:
                text += self.competence[key]
            cell.text = text

            self.change_map(table_num=table_num, offset=1, fromkey='2')


    def change_seven_ten(self):

        table_num, row_num = self.docx_map['7.3.2']
        table = self.document.tables[table_num]

        self.delete_between_rows((table_num, row_num),
                                 (table_num,len(self.document.tables[table_num].rows))
                                 )

        for _ in range(table_num+1, len(self.document.tables)):
            self.delete_table(self.document.tables[table_num+1])


        #Читаем нужную таблицу
        if self.learning_plan_info['learning_form'] and self.learning_plan_info['kurs']:
            doc = docx.Document(os.path.join(BASE_DIR,'data','Очка_к.docx'))
            tbl_to_add = copy.deepcopy(doc.tables[0]._tbl)
        elif self.learning_plan_info['learning_form'] and not self.learning_plan_info['kurs']:
            doc = docx.Document(os.path.join(BASE_DIR, 'data', 'Очка_бк.docx'))
            tbl_to_add = copy.deepcopy(doc.tables[0]._tbl)
        elif not self.learning_plan_info['learning_form'] and self.learning_plan_info['kurs']:
            doc = docx.Document(os.path.join(BASE_DIR, 'data', 'Заочка_к.docx'))
            tbl_to_add = copy.deepcopy(doc.tables[0]._tbl)
        elif not self.learning_plan_info['learning_form'] and not self.learning_plan_info['kurs']:
            doc = docx.Document(os.path.join(BASE_DIR, 'data', 'Заочка_бк.docx'))
            tbl_to_add = copy.deepcopy(doc.tables[0]._tbl)

        self.document.tables[-1]._tbl.addnext(tbl_to_add)

        for row in self.document.tables[-1].rows:
            self.set_cell_border(
                row.cells[0],
                top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                end={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
            )

            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            cell = row.cells[0]
            # cell.paragraphs[0].runs[0].style.font.name = 'Times New Roman'
            # cell.paragraphs[0].runs[0].style.font.size = docx.shared.Pt(19)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            cell.paragraphs[0].paragraph_format.space_after = Cm(0)




    def change_seven(self):
        '''
        Добавляем правильные электронные каталоги в раздел 7.3.2
        :return:
        '''
        chapter = self.rpd_chapters.get_chapter('7.3.2')['Каталоги']
        table_num, row_num = self.docx_map['7.3.2']
        table = self.document.tables[table_num]

        #Удаляем лишние ячейки
        rows_between_this_table, rows_between_next_table = self.delete_between_rows(self.docx_map['7.3.2'], self.docx_map['8'])



        #добавляем ячейки с каталогами
        for i in range(len(chapter)):
            self.add_row(table.rows[row_num], table.rows[row_num])

        for i, chapter_part in enumerate(chapter):
            cell = table.rows[row_num+i+1].cells[0]
            cell.text = chapter_part
            table.rows[row_num+i+1].height_rule = WD_ROW_HEIGHT_RULE.AUTO

            #cell.paragraphs[0].runs[0].style.font.name = 'Times New Roman'
            #cell.paragraphs[0].runs[0].style.font.size = docx.shared.Pt(19)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            cell.paragraphs[0].paragraph_format.space_after = Cm(0)

        # Меняем название раздела на правильное
        table.rows[row_num].cells[0].paragraphs[0].text = "7.3.2. Профессиональные базы данных и информационно-справочные системы"
        table.rows[row_num].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[row_num].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # Высчитываем насколько надо поправить карту разделов
        offset = len(chapter) - rows_between_this_table
        self.change_map(table_num=table_num, offset=offset, fromkey = '7.3.2')

        if rows_between_next_table !=0:
            self.change_map(table_num=table_num+1, offset=-rows_between_next_table, fromkey='7.3.2')


    def change_eight(self):
        '''
        Добавляем правильный текст в раздел 8 с учетом формы обучения и курсовой
        learning_form = True - очная
        Kurs = True - есть курсовая
        :return:
        '''

        table_num, row_num = self.docx_map['8']
        table = self.document.tables[table_num]

        #Удаляем лишние ячейки
        rows_between_this_table, rows_between_next_table = self.delete_between_rows(self.docx_map['8'], self.docx_map['9'])


        # добавляем ячейку куда будем вносить данные
        #cp_table, cp_row = self.docx_map['1']
        row = table.rows[row_num]
        self.add_row(table.rows[row_num], row)


        chapter = self.rpd_chapters.get_chapter('8')

        if self.learning_plan_info['learning_form'] and self.learning_plan_info['kurs']:
            chapter = chapter['Очка_К']
        elif self.learning_plan_info['learning_form'] and not self.learning_plan_info['kurs']:
            chapter = chapter['Очка']
        elif not self.learning_plan_info['learning_form'] and self.learning_plan_info['kurs']:
            chapter = chapter['Заочка_К']
        elif not self.learning_plan_info['learning_form'] and not self.learning_plan_info['kurs']:
            chapter = chapter['Заочка']

        table.rows[row_num+1].cells[0].text = chapter
        table.rows[row_num + 1].height_rule = WD_ROW_HEIGHT_RULE.AUTO
        self.set_white_color(table.rows[row_num + 1])

        # Меняем название раздела на правильное
        table.rows[row_num].cells[0].paragraphs[
            0].text = SECTIONS['8']
        table.rows[row_num].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[row_num].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Высчитываем насколько надо поправить карту разделов
        offset = len(chapter) - rows_between_this_table
        self.change_map(table_num=table_num, offset=offset, fromkey = '8')

        if rows_between_next_table != 0:
            self.change_map(table_num=table_num+1, offset=-rows_between_next_table, fromkey='8')

    def change_nine(self):
        '''
        Добавляем правильный текст в раздел 9
        :return:
        '''

        table_num, row_num = self.docx_map['9']
        table = self.document.tables[table_num]

        # Удаляем лишние ячейки
        rows_between_this_table, rows_between_next_table = self.delete_between_rows(self.docx_map['9'], self.docx_map['10'])


        # добавляем ячейку куда будем вносить данные
        #cp_table, cp_row = self.docx_map['1']
        row = table.rows[row_num]
        self.add_row(table.rows[row_num], row)

        chapter = self.rpd_chapters.get_chapter('9')
        if self.learning_plan_info['kurs']:
            chapter = chapter['С_курсовой']
        else:
            chapter = chapter['Без_курсовой']

        table.rows[row_num + 1].cells[0].text = chapter
        table.rows[row_num + 1].height_rule = WD_ROW_HEIGHT_RULE.AUTO
        self.set_white_color(table.rows[row_num + 1])

        # Меняем название раздела на правильное
        table.rows[row_num].cells[0].paragraphs[
            0].text = SECTIONS['9']
        table.rows[row_num].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[row_num].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Высчитываем насколько надо поправить карту разделов
        offset = len(chapter) - rows_between_this_table
        self.change_map(table_num=table_num, offset=offset, fromkey='9')

        if rows_between_next_table !=0 :
            self.change_map(table_num=table_num+1, offset=-rows_between_next_table, fromkey='9')

    def change_ten(self):
        '''
        Правим 10 раздел
        :return:
        '''

        table_num, row_num = self.docx_map['10']
        table = self.document.tables[table_num]


        # Удаляем лишние ячейки
        rows_between_this_table, rows_between_next_table = self.delete_between_rows((table_num, row_num+1), (table_num,len(table.rows)))

        row = table.rows[row_num]
        self.add_row(table.rows[row_num], row)

        chapter = self.rpd_chapters.get_chapter('10')
        chapter = chapter['Текст']

        table.rows[row_num + 1].cells[0].text = chapter
        table.rows[row_num + 1].height_rule = WD_ROW_HEIGHT_RULE.AUTO
        self.set_white_color(table.rows[row_num + 1])

        # Меняем название раздела на правильное
        table.rows[row_num].cells[0].paragraphs[
            0].text = SECTIONS['10']
        table.rows[row_num].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[row_num].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER




    def delete_table(self,table):
        '''
        Удаляем таблицу
        :param table:
        :return:
        '''

        tb = table._element
        tb.getparent().remove(tb)


    def change_map(self,**kwargs):
        '''
        Сдвиг карты разделов, после добавления строк в таблицы
        :param kwargs:
        :return:
        '''

        keys = list(self.docx_map.keys())
        keys = keys[keys.index(kwargs['fromkey'])+1:]

        for key in keys:

            table_num, row_num = self.docx_map[key]

            if table_num == kwargs['table_num']:
                self.docx_map[key] = (table_num, row_num+kwargs['offset'])


    def delete_between_rows(self, from_pair, till_pair):

        table_from, row_from = from_pair
        table_till, row_till = till_pair


        table = self.document.tables[table_from]

        if table_from == table_till:
            rows_between_this_table = row_till - row_from - 1
            for _ in range(rows_between_this_table):
                self.remove_row(table, table.rows[row_from + 1])
            rows_between_next_table = 0
        else:
            rows_between_this_table = len(table.rows) - row_from - 2
            for _ in range(rows_between_this_table):
                self.remove_row(table, table.rows[row_from + 1])
            rows_between_next_table = row_till
            for _ in range(rows_between_next_table):
                self.remove_row(self.document.tables[table_till], self.document.tables[table_till].rows[0])

        return rows_between_this_table, rows_between_next_table


    @staticmethod
    def set_white_color(row):
        for cell in row.cells:
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm_2)

    @staticmethod
    def add_row(row_insert_after, row_adding):
        '''
        Добавить строку в таблице
        :param row:
        :return:
        '''
        tr = copy.deepcopy(row_adding._tr)
        row_insert_after._tr.addnext(tr)


    @staticmethod
    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def set_cell_border(self, cell: _Cell, **kwargs):
        """
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))


    def __del__(self):

        self.document.save(os.path.join(self.path_to_save, self.filename))

def main():
    starting_dir = r'D:\RPD_TEST\RPD8'
    saving_dir = r'D:\RPD_TEST\MODIFIED_SECTION'
    for dir_l_1 in os.listdir(starting_dir):
        print (f"\n\nОбрабатываю год {dir_l_1}\n\n")
        for dir_l_2 in os.listdir(os.path.join(starting_dir,dir_l_1)):
            print(f"\nОбрабатываю направление подготовки {dir_l_2}\n")
            medium_dir = os.path.join(dir_l_1, dir_l_2)
            path_to_save = os.path.join(saving_dir,medium_dir)
            path_to_read = os.path.join(starting_dir, medium_dir)
            if not os.path.isdir(path_to_save):
                os.makedirs(path_to_save)

            if os.path.exists(os.path.join(path_to_save, "log.txt")):
                os.remove(os.path.join(path_to_save, "log.txt"))
            with open(os.path.join(path_to_save, "log.txt"), 'x'):
                pass
            for file in os.listdir(path_to_read):
            #path = r'D:\RPD_TEST\TEMP'
            #file = 'Алгоритмы и анализ сложности.docx'

                print(f"Обрабатываю файл {file}")
                try:
                    sections = Sections_changer(path_to_read, path_to_save, file)
                    sections.change_one()
                    sections.change_seven_ten()
                    del sections
                except Exception as e:
                    print(f"Error found with file: {file}. Error type - {e}")
                    with open(os.path.join(path_to_save,'log.txt'), 'a') as f:
                        f.write(f"Error with file {file}\n")
    #sections.change_seven()
    #sections.change_eight()
    #sections.change_nine()
    #sections.change_ten()


if __name__ == '__main__':
    main()


