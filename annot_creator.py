from sections_changer import RPD_CHAPTERS
import os
import mammoth
import re
import pandas as pd
import tqdm
import docx
import numpy as np
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Inches, Mm
from docx.table import _Cell
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy
from docxtpl import DocxTemplate
from PyPDF2 import PdfFileReader, PdfFileWriter
from docx2pdf import convert




SECTION_1 = re.compile(r'1. ЦЕЛИ ОСВОЕНИЯ ДИСЦИПЛИНЫ|1. ОБЩИЕ ПОЛОЖЕНИЯ ПО ПРАКТИКЕ')
SECTION_2 = re.compile(r'2. МЕСТО ДИСЦИПЛИНЫ В СТРУКТУРЕ ОПОП|2. МЕСТО ПРАКТИКИ В СТРУКТУРЕ ОБРАЗОВАТЕЛЬНОЙ ПРОГРАММЫ')
SECTION_3 = re.compile(r'3. КОМПЕТЕНЦИИ ОБУЧАЮЩЕГОСЯ, ФОРМИРУЕМЫЕ В РЕЗУЛЬТАТЕ ОСВОЕНИЯ ДИСЦИПЛИНЫ \(МОДУЛЯ\)|3. ПЕРЕЧЕНЬ ПЛАНИРУЕМЫХ РЕЗУЛЬТАТОВ ОБУЧЕНИЯ')
SECTION_4 = re.compile(r'4. ОБЪЕМ ДИСЦИПЛИНЫ \(МОДУЛЯ\)|4. ОБЪЕМ ПРАКТИКИ')
SECTION_5 = re.compile(r'5. СТРУКТУРА И СОДЕРЖАНИЕ ДИСЦИПЛИНЫ|5. СОДЕРЖАНИЕ ПРАКТИКИ')
SECTION_3_1 = re.compile(r'В результате освоения дисциплины обучающийся должен|По окончанию практики обучающийся должен')


LEARNING_QUALIFICATION_REGEX = re.compile(r'бакалавр|магистр')
LEARNING_YEAR_REGEX = re.compile(r'Челябинск (\d{4}) г.')
LEARNING_FORM_REGEX = re.compile(r'(очная|заочная|очно-заочная)')
LEARNING_DISC_NAME_REGEX = re.compile(r'Рабочая программа дисциплины \(модуля\)\*\n\n|Рабочая программа практики\*\n\n')
LEARNING_DISC_REGEX = re.compile(r'Рабочая программа дисциплины \(модуля\)\*\n\n')

LEARNING_CODE_REGEX = re.compile(r'Направление подготовки \(специальность\)\n\n')
LEARNING_PROFILE_REGEX = re.compile(r'Направленность \(профиль\)\n\n')


STOP_WORDS = (
            '© ФГБОУ ВО «ЧелГУ»',
            'Рабочая программа дисциплины "',
            'Рабочая программа практики "',
            'стр. ',
            )



def create_header(document, context):
    '''
    Создаем верхний колонтитул
    :param document:
    :return:
    '''
    section = document.sections[0]
    header = section.header

    first_header = docx.Document(os.path.join('headers', 'first_header.docx'))
    cell = first_header.tables[1].rows[0].cells[0]
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.paragraphs[0].style.font.name = 'Times New Roman'
    cell.paragraphs[0].style.font.size = docx.shared.Pt(9)

    if context['discipline']:
        header_text = "Рабочая программа дисциплины "
    else:
        header_text = "Рабочая программа практики "

    header_text += f"\"{context['discipline_name']}\" по направлению подготовки (специальности) "
    header_text += f"\"{context['learning_code_short']}\" направленности (профилю) "
    header_text += f"{context['learning_profile']} ФГБОУ ВО «ЧелГУ»"
    cell.paragraphs[0].text = header_text

    header_table = copy.deepcopy(first_header.tables[1]._tbl)
    header.paragraphs[0]._p.addnext(header_table)

    # document.save(os.path.join(RPD_MODIFIED_DIR, self.filepath, self.filename_to_save))

    # document = docx.Document(os.path.join(RPD_MODIFIED_DIR, self.filepath, self.filename_to_save))
    return document


def create_footer(footer):
    '''
    Создаем футер
    :param self:
    :param footer:
    :return:
    '''
    footer.paragraphs[0].text = "© ФГБОУ ВО «ЧелГУ»"
    footer.paragraphs[0].style.font.name = 'Times New Roman'
    footer.paragraphs[0].style.font.size = docx.shared.Pt(9.5)

def first_section_creator(first_section, practice=False):
    '''
    Преобразуем в список содержимое 1 секции
    :param first_section:
    :return:
    '''
    if practice:
        structure = ['1. ОБЩИЕ ПОЛОЖЕНИЯ ПО ПРАКТИКЕ']
    else:
        structure = ['1. ЦЕЛИ ОСВОЕНИЯ ДИСЦИПЛИНЫ']
    for comp in ['УК', 'ОПК', 'ПК']:
        first_section = re.sub('ОПК', '\nОПК', first_section)
    first_section = first_section.strip()
    first_section = re.sub('[\n]{2,}', '\n', first_section)
    structure.extend(first_section.split('\n'))
    structure = np.array(structure).reshape((-1,1))
    return structure



def second_section_creator(second_section, practice=False):
    '''
    Преобразуем 2 секцию в нужный формат
    :param second_section:
    :return:
    '''

    if practice:
        structure = np.array([['2. МЕСТО ПРАКТИКИ В СТРУКТУРЕ ОБРАЗОВАТЕЛЬНОЙ ПРОГРАММЫ','']])
    else:
        structure = np.array([['2. МЕСТО ДИСЦИПЛИНЫ В СТРУКТУРЕ ОПОП', '']])
    second_section = second_section.strip()
    second_section = re.sub('[\n]{2,}', '\n', second_section)
    structure = np.vstack([structure, second_section.split('\n')[:2]])
    return structure


def third_section_creator(third_section, practice=False):
    '''
    Преобразуем 3 раздел в нужный формат
    :param third_section:
    :param practice:
    :return:
    '''

    if practice:
        structure = ['3. ПЕРЕЧЕНЬ ПЛАНИРУЕМЫХ РЕЗУЛЬТАТОВ ОБУЧЕНИЯ']
    else:
        structure = ['3. КОМПЕТЕНЦИИ ОБУЧАЮЩЕГОСЯ, ФОРМИРУЕМЫЕ В РЕЗУЛЬТАТЕ ОСВОЕНИЯ ДИСЦИПЛИНЫ (МОДУЛЯ)']

    third_section = third_section.strip()
    third_section = re.sub('[\n]{2,}', '\n', third_section)
    structure.extend(third_section.split('\n'))
    structure = np.array(structure).reshape((-1,1))
    return structure


def forth_section_creator(forth_section, practice=False):
    '''
    Преобразуем 4 раздел в нужный формат
    :param forth_section:
    :param practice:
    :return:
    '''

    if practice:
        structure = np.array([['4. ОБЪЕМ ПРАКТИКИ', '']])
    else:
        structure = np.array([['4. ОБЪЕМ ДИСЦИПЛИНЫ (МОДУЛЯ)', '']])
    forth_section = forth_section.strip()
    forth_section = re.sub('[\n]{2,}', '\n', forth_section)
    forth_section = forth_section.split('\n')
    forth_section[7] = forth_section[7] + '\n'
    structure = np.vstack([structure,[forth_section[0],forth_section[1]]])
    structure = np.vstack([structure,['\n'.join(forth_section[2:7]),'\n'.join(forth_section[7:])]])
    return structure


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    if kwargs.get('colored'):
        shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
        tc.get_or_add_tcPr().append(shading_elm)

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_second_page_table(sections, context):
    '''
    Создаем 2 страницу аннотации
    :param sections:
    :return:
    '''

    document = docx.Document()
    section = document.sections[0]
    section.left_margin = Mm(20.0)
    section.right_margin = Mm(10.0)
    section.top_margin = Mm(0.1)
    section.bottom_margin = Mm(9.5)
    section.header_distance = Mm(8.0)
    section.footer_distance = Mm(12.5)



    for sec_num, section in enumerate(sections):
        table = document.add_table(rows=section.shape[0], cols=section.shape[1])
        for row_n, row in enumerate(section):
            if row_n == 0:
                table.rows[row_n].cells[0].merge(table.rows[row_n].cells[-1])
                cell = table.rows[row_n].cells[0]
                cell.text = row[0]
                table.rows[row_n].height_rule = WD_ROW_HEIGHT_RULE.AUTO
                cell.paragraphs[0].style.font.name = 'Times New Roman'
                cell.paragraphs[0].style.font.size = docx.shared.Pt(9)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Cm(0.1)
                set_cell_border(
                    cell,
                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    colored=True,
                )
                table.rows[row_n].cells[0].paragraphs[0].runs[0].font.bold = True
                table.rows[row_n].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                continue

            for col_n, col in enumerate(row):
                cell = table.rows[row_n].cells[col_n]
                cell.text = col
                table.rows[row_n].height_rule = WD_ROW_HEIGHT_RULE.AUTO
                cell.paragraphs[0].style.font.name = 'Times New Roman'
                cell.paragraphs[0].style.font.size = docx.shared.Pt(9)
                cell.paragraphs[0].paragraph_format.line_spacing = 2.0
                cell.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Cm(0.02)
                cell.paragraphs[0].paragraph_format.space_before = Cm(0.05)
                set_cell_border(
                    cell,
                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                )

                #На случай компетенцйй надо выделить жирным
                if any(compet in col.lower() for compet in ['пк', 'ук', 'опк', 'знать', 'уметь', 'владеть']) and sec_num>0:
                    cell.paragraphs[0].runs[0].font.bold = True
        res = table.add_row().cells

    document = create_header(document, context) ### создаем колонтитул 2 страницы
    create_footer(document.sections[0].footer) ### создаем футер 2 страницы
    document.save(r'temp\body.docx')



def create_first_page(context):
    '''
    Генерация первой страницы аннотации
    :return:
    '''

    doc = DocxTemplate(os.path.join('headers', 'first_page_template.docx'))

    '''context = {
        'annot_type': "Аннотация рабочей программы дисциплины (модуля)",
        'discipline_name': "Базы и хранилища данных",
        'learning_code': "02.03.02 Фундаментальная информатика и информационные технологии",
        'learning_profile': "Инженерия программного обеспечения",
        'qualification': "бакалавр",
        'learning_form': "очно-заочная",
        'learning_code_short': "Фундаментальная информатика и информационные технологии",
        'learning_from_year': "2021",

    }'''

    doc.render(context)
    doc.save(os.path.join('temp', 'first_page.docx'))

def create_document_info(text):
    '''
    Создаем файл контекста
    :param text:
    :return:
    '''
    context = {}

    context['learning_from_year'] = LEARNING_YEAR_REGEX.findall(text)[0]
    context['learning_form'] = LEARNING_FORM_REGEX.findall(text)[0]
    context['qualification'] = LEARNING_QUALIFICATION_REGEX.findall(text.lower())[0]
    context['discipline_name'] = LEARNING_DISC_NAME_REGEX.split(text)[1].strip().split('\n')[0]
    context['learning_code'] = LEARNING_CODE_REGEX.split(text)[1].strip().split('\n')[0]
    context['learning_profile'] = LEARNING_PROFILE_REGEX.split(text)[1].strip().split('\n')[0]
    context['learning_code_short'] = context['learning_code'].split(' ')[1]
    if len(LEARNING_DISC_REGEX.findall(text))>0:
        context['annot_type'] = 'Аннотация рабочей программы дисциплины (модуля)'
        context['discipline'] = True
    else:
        context['annot_type'] = 'Аннотация рабочей программы практики'
        context['discipline'] = False


    return context


def create_pdf(filename = '.\\temp\\result.pdf'):
    '''
    Генерируем PDF файл соединяя 1 и 2 страницы
    :param filename:
    :return:
    '''

    convert('.\\temp\\body.docx', '.\\temp\\body.pdf')
    convert('.\\temp\\first_page.docx', '.\\temp\\first_page.pdf')
    title = PdfFileReader(os.path.join('temp', 'first_page.pdf'))
    body = PdfFileReader(os.path.join('temp', 'body.pdf'))
    # body = PdfFileReader(body_filename)
    output = PdfFileWriter()
    for page_number in range(title.getNumPages()):
        output.addPage(title.getPage(page_number))
    for page_number in range(body.getNumPages()):
        output.addPage(body.getPage(page_number))

    with open(filename, 'wb') as f:
        output.write(f)

    return 0

def create_clear_dir(path):
    '''
    Функция очистки или создания указанного пути
    :param path:
    :return:
    '''

    print(f"### Начинаю создание или очистку каталогов по пути {path} ###")

    if not os.path.exists(path):
        os.mkdir(path)
    else:
        for root, dirs, files in os.walk(path):
            for f in files:
                os.unlink(os.path.join(root, f))




def annot_control(path, path_to_save):
    '''
    Метод по созданию аннотаций
    :return:
    '''

    print(f"### Начинаю создание аннотаций на основе РПД из папки {path} ###")

    files =  [os.path.join(path, file) for file in os.listdir(path) if file.endswith('.docx')]
    #files =  [os.path.join(path, file) for file in os.listdir(path) if "Создание бизнес-приложений в системе" in file]

    create_clear_dir(path_to_save)




    for file in files:

        print(f"*** Начинаю создание аннотации файла {os.path.split(file)[1]} ***")

        with open(file, 'rb') as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value  # The raw text
            messages = result.messages  # Any messages



        context = create_document_info(text) ### создаем контекст
        create_first_page(context) ### создаем 1 страницу аннотации

        text_for_sections = [txt for txt in text.strip().split('\n')
                             if all(stop_sentence not in txt for stop_sentence in STOP_WORDS)] ### Создаем текст для обработки по секциям
        text_for_sections = '\n'.join(text_for_sections)

        ### Создаем разделы в виде текста


        first_section = SECTION_2.split(SECTION_1.split(text_for_sections)[1])[0]
        second_section = SECTION_3.split(SECTION_2.split(text_for_sections)[1])[0]
        third_section = SECTION_3_1.split(SECTION_4.split(SECTION_3.split(text_for_sections)[1])[0])[0]
        forth_section = SECTION_5.split(SECTION_4.split(text_for_sections)[1])[0]

        ### Создаем секции 2 страницы РПД

        sections = []
        sections.append(first_section_creator(first_section))
        sections.append(second_section_creator(second_section))
        sections.append(third_section_creator(third_section))
        sections.append(forth_section_creator(forth_section))
        create_second_page_table(sections, context) ### Создаем 2 страницу аннотации


        ### Создаем аннотацию как PDF файл
        annot_filename = f"{context['discipline_name']}.pdf"
        create_pdf(filename=os.path.join(path_to_save, annot_filename))

        print(f"*** Аннотация файла {os.path.split(file)[1]} успешно создана и сохранена в файл {annot_filename} ***\n")



    print(f"### Успешно закончил создание аннотаций из пути {path} ###")




PATH_TO_RPD = r"D:\RPD_TEST\RPD_body"
PATH_TO_SAVE = r"D:\RPD_TEST\ANNOT"

paths = [(os.path.join(PATH_TO_RPD, folder), os.path.join(PATH_TO_SAVE, folder)) for folder in os.listdir(PATH_TO_RPD)
         if os.path.isdir(os.path.join(PATH_TO_RPD, folder))]

for path in paths[7:]:
    annot_control(path[0], path[1])
